from pathlib import Path
import re
import zipfile

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN = ROOT / "materials" / "scripts" / "AIDK_W1_Speaker_Script.md"
DOCX = ROOT / "materials" / "scripts" / "AIDK_W1_Facilitator_Script_2026-07-15.docx"


markdown = MARKDOWN.read_text(encoding="utf-8")
slide_matches = list(re.finditer(r"(?m)^## Slide (\d+): (.+)$", markdown))
assert len(slide_matches) == 45, f"Expected 45 Markdown slide sections, found {len(slide_matches)}"
assert [int(match.group(1)) for match in slide_matches] == list(range(1, 46))

exercise_slides = {17, 19, 23, 25, 27, 29, 31, 33, 35, 37, 39, 40, 41}
for position, match in enumerate(slide_matches):
    slide_number = int(match.group(1))
    if slide_number not in exercise_slides:
        continue
    end = slide_matches[position + 1].start() if position + 1 < len(slide_matches) else len(markdown)
    section = markdown[match.start():end]
    assert "**Facilitator Check:**" in section, f"Slide {slide_number} is missing a facilitator answer"

with zipfile.ZipFile(DOCX) as archive:
    bad_member = archive.testzip()
    assert bad_member is None, f"Corrupt DOCX member: {bad_member}"
    names = set(archive.namelist())
    for required in {
        "[Content_Types].xml",
        "word/document.xml",
        "word/styles.xml",
        "word/header1.xml",
        "word/footer1.xml",
    }:
        assert required in names, f"Missing required DOCX part: {required}"

document = Document(DOCX)
docx_slide_headings = [
    paragraph.text
    for paragraph in document.paragraphs
    if paragraph.style.name == "Heading 1" and paragraph.text.startswith("Slide ")
]
assert len(docx_slide_headings) == 45, f"Expected 45 DOCX slide headings, found {len(docx_slide_headings)}"
assert [int(re.match(r"Slide (\d+):", heading).group(1)) for heading in docx_slide_headings] == list(range(1, 46))

code_blocks = [paragraph for paragraph in document.paragraphs if paragraph.style.name == "Code Block"]
assert len(code_blocks) >= 25, f"Expected at least 25 code blocks, found {len(code_blocks)}"
assert len(document.tables) == 1, f"Expected one schedule table, found {len(document.tables)}"

section = document.sections[0]
assert abs(section.page_width - Inches(8.5)) < 100
assert abs(section.page_height - Inches(11)) < 100
for margin in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin):
    assert abs(margin - Inches(1)) < 100

normal = document.styles["Normal"]
heading_one = document.styles["Heading 1"]
assert normal.font.name == "Calibri"
assert normal.font.size == Pt(11)
assert heading_one.font.size == Pt(16)
assert document.core_properties.author == "SPAI"
assert document.core_properties.title.startswith("AI Don't Know Workshop 1")

header_text = " ".join(paragraph.text for paragraph in section.header.paragraphs)
footer_text = " ".join(paragraph.text for paragraph in section.footer.paragraphs)
assert "SPAI" in header_text and "Workshop 1" in header_text
assert "Page" in footer_text

print(
    "Speaker script DOCX check passed: "
    f"45 slide headings, {len(code_blocks)} code blocks, "
    "13 exercise answer keys, valid package, Letter geometry, and expected styles."
)
