from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "materials" / "scripts" / "AIDK_W1_Speaker_Script.md"
OUTPUT = ROOT / "materials" / "scripts" / "AIDK_W1_Facilitator_Script_2026-07-15.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x20, 0x37, 0x48)
COPPER = RGBColor(0xB7, 0x5F, 0x45)
AMBER = RGBColor(0xA7, 0x68, 0x17)
GRAY = RGBColor(0x5F, 0x63, 0x68)
MUTED = RGBColor(0x77, 0x77, 0x77)
BLACK = RGBColor(0x18, 0x18, 0x18)
CODE_FILL = "F2F4F7"
CALLOUT_FILL = "FFF4E8"
BORDER = "D9E2EC"


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run_node = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run_node.append(text)
    fld.append(run_node)
    paragraph._p.append(fld)


def add_inline_markdown(paragraph, text, base_size=11, base_color=BLACK):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=base_size - 0.5, color=DARK_BLUE)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=base_color)


def add_code_block(doc, code):
    paragraph = doc.add_paragraph(style="Code Block")
    run = paragraph.add_run(code.rstrip())
    set_run_font(run, name="Consolas", size=9.2, color=RGBColor(0x20, 0x25, 0x2A))
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), BORDER)
        node.set(qn("w:space"), "5")
        borders.append(node)
    p_pr.append(borders)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = DARK_BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    label = styles.add_style("Script Label", WD_STYLE_TYPE.PARAGRAPH)
    label.base_style = normal
    label.font.name = "Calibri"
    label.font.size = Pt(9.5)
    label.font.bold = True
    label.font.color.rgb = COPPER
    label.paragraph_format.space_before = Pt(8)
    label.paragraph_format.space_after = Pt(2)
    label.paragraph_format.keep_with_next = True

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = normal
    code.font.name = "Consolas"
    code.font.size = Pt(9.2)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    code.paragraph_format.keep_together = True


def add_cover(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("SPAI | AI Don't Know Workshop 1")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("FACILITATOR SCRIPT")
    set_run_font(run, size=10, color=COPPER, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    run = title.add_run("AI Don't Know Workshop 1")
    set_run_font(run, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Python Foundations - Slide-by-Slide Run of Show")
    set_run_font(run, size=13.2, color=GRAY)

    for label, value in (
        ("Workshop", "15 July 2026 | 3:00 PM-5:00 PM"),
        ("Notebook", "01_AIDK_W1_Starter.ipynb"),
        ("Slides", "spai-aidk-workshop1.vercel.app"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=BLACK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    schedule = [
        ("3:00", "Welcome and setup"),
        ("3:20", "Python and JavaScript"),
        ("3:44", "Python basics"),
        ("4:28", "Combined practice"),
    ]
    for cell, (time, label) in zip(table.rows[0].cells, schedule):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=120, start=110, bottom=120, end=110)
        shade_cell(cell, CALLOUT_FILL)
        set_cell_border(cell, color="E8D7C1", size="4")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        time_run = p.add_run(time)
        set_run_font(time_run, size=12, color=AMBER, bold=True)
        p.add_run().add_break()
        label_run = p.add_run(label)
        set_run_font(label_run, size=8.8, color=BLACK, bold=True)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Presenter principle: explain one idea, run the matching example, then give participants time to type the prompt themselves.")
    set_run_font(run, size=11, color=DARK_BLUE, bold=True)

    doc.add_page_break()


def build_document():
    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    start = next(i for i, line in enumerate(lines) if line == "## How to Use This Script")

    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    page_break_slides = {1, 8, 14, 22, 38, 42}
    code_buffer = []
    in_code = False
    index = start

    while index < len(lines):
        line = lines[index]

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                add_code_block(doc, "\n".join(code_buffer))
                in_code = False
            index += 1
            continue

        if in_code:
            code_buffer.append(line)
            index += 1
            continue

        if not line.strip():
            index += 1
            continue

        if line.startswith("## "):
            heading = line[3:]
            slide_match = re.match(r"Slide (\d+):", heading)
            if slide_match and int(slide_match.group(1)) in page_break_slides and int(slide_match.group(1)) != 1:
                doc.add_page_break()
            p = doc.add_paragraph(heading, style="Heading 1")
            if slide_match:
                p.paragraph_format.page_break_before = False
            index += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, line[2:])
            index += 1
            continue

        label_match = re.match(r"^\*\*(.+?):\*\*\s*(.*?)(?:\s{2})?$", line)
        if label_match:
            label, value = label_match.groups()
            if label in {"Timing", "Speaker"} and value:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                label_run = p.add_run(f"{label}: ")
                set_run_font(label_run, size=9.5, color=COPPER, bold=True)
                value_run = p.add_run(value)
                set_run_font(value_run, size=9.5, color=GRAY, italic=True)
            else:
                p = doc.add_paragraph(style="Script Label")
                p.add_run(label.upper())
                if value:
                    value_run = p.add_run(f"  {value}")
                    set_run_font(value_run, size=10.5, color=BLACK)
            index += 1
            continue

        p = doc.add_paragraph()
        if line in {"Expected output:", "One valid answer is:"}:
            run = p.add_run(line)
            set_run_font(run, size=11, color=BLACK, bold=True)
        else:
            add_inline_markdown(p, line)
        index += 1

    props = doc.core_properties
    props.title = "AI Don't Know Workshop 1: Python Foundations - Facilitator Script"
    props.subject = "Slide-by-slide facilitator script and notebook run of show"
    props.author = "SPAI"
    props.keywords = "SPAI, Python, workshop, facilitator script, AI Don't Know"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build_document()
